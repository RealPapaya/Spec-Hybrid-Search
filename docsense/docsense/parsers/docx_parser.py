"""
DOCX parser — python-docx.

Extracts paragraphs, headings (detected by style name), and table content.
Heading blocks are tagged with ``block_type="heading"`` so the chunking
engine can use them as natural section boundaries.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table

from docsense.parsers.base import ParsedDocument, ParserRegistry, TextBlock

logger = logging.getLogger(__name__)


class DocxParser:
    """Parse ``.docx`` files into structured :class:`TextBlock` objects."""

    def parse(self, filepath: Path) -> ParsedDocument:
        """
        Extract text blocks from a Word document.

        Parameters
        ----------
        filepath:
            Path to the ``.docx`` file.

        Returns
        -------
        ParsedDocument
            Blocks in document order; headings are tagged separately.
        """
        blocks: list[TextBlock] = []
        current_section: str | None = None

        try:
            doc = DocxDocument(str(filepath))

            for element in doc.element.body:
                # Resolve the local tag name (strip namespace URI)
                tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

                if tag == "p":
                    para = self._find_paragraph(doc, element)
                    if para is None:
                        continue

                    text = para.text.strip()
                    if not text:
                        continue

                    style_name = (para.style.name or "").lower()

                    if "heading" in style_name:
                        current_section = text
                        blocks.append(TextBlock(
                            text=text,
                            section_title=current_section,
                            block_type="heading",
                        ))
                    else:
                        blocks.append(TextBlock(
                            text=text,
                            section_title=current_section,
                            block_type="paragraph",
                        ))

                elif tag == "tbl":
                    table = self._find_table(doc, element)
                    if table is None:
                        continue
                    table_text = self._table_to_text(table)
                    if table_text.strip():
                        blocks.append(TextBlock(
                            text=table_text,
                            section_title=current_section,
                            block_type="table",
                        ))

        except Exception as exc:
            logger.error("DOCX parse error for %s: %s", filepath, exc)
            return ParsedDocument(
                filename=filepath.name,
                file_type=".docx",
                error=str(exc),
            )

        return ParsedDocument(
            filename=filepath.name,
            file_type=".docx",
            blocks=blocks,
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _find_paragraph(doc: DocxDocument, element):
        """Return the Paragraph object whose underlying XML element matches."""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    @staticmethod
    def _find_table(doc: DocxDocument, element):
        """Return the Table object whose underlying XML element matches."""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    @staticmethod
    def _table_to_text(table: Table) -> str:
        """Render a DOCX table as pipe-separated rows."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


# Register with the global parser registry
ParserRegistry.register(".docx", DocxParser())
